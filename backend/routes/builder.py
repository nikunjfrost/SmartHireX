import io
from flask import Blueprint, request, jsonify, send_file
from docx import Document
import os
import google.generativeai as genai
from dotenv import load_dotenv
from services.dataset_loader import JOB_ROLES

builder_bp = Blueprint('builder', __name__)

@builder_bp.route('/optimize', methods=['POST'])
def optimize_resume():
    data = request.get_json()
    if not data:
        return jsonify({"error": "No data provided"}), 400
        
    personal_info = data.get('personalInfo', {})
    summary = data.get('summary', '')
    experience = data.get('experience', [])
    education = data.get('education', [])
    skills = data.get('skills', [])
    desired_role = data.get('desiredRole', '').strip()
    
    # 1. Optimize data based on desired_role
    optimized_skills = list(skills)
    optimized_summary = summary
    
    if desired_role:
        # Find closest match or exact match in JOB_ROLES
        role_key = None
        for role_name in JOB_ROLES.keys():
            if desired_role.lower() in role_name.lower() or role_name.lower() in desired_role.lower():
                role_key = role_name
                break
                
        if role_key:
            role_info = JOB_ROLES[role_key]
            required_skills = role_info.get("required_skills", [])
            
            # Add missing required skills
            existing_skills_lower = [s.lower() for s in optimized_skills]
            for req_skill in required_skills:
                if req_skill.lower() not in existing_skills_lower:
                    optimized_skills.append(req_skill)
                    
            # Enhance summary
            if not optimized_summary:
                optimized_summary = f"Results-driven {role_key} with expertise in {', '.join(required_skills[:3])}. Proven ability to deliver high-quality solutions and drive business success."
            else:
                optimized_summary = f"Targeting {role_key} positions. " + optimized_summary
        elif not optimized_summary:
            optimized_summary = f"Dedicated professional seeking {desired_role} opportunities."

    # 2. Generate DOCX
    doc = Document()
    
    # Personal Info Header
    name = f"{personal_info.get('firstName', '')} {personal_info.get('lastName', '')}".strip()
    if name:
        doc.add_heading(name, 0)
        
    contact_parts = []
    if personal_info.get('email'): contact_parts.append(personal_info.get('email'))
    if personal_info.get('phone'): contact_parts.append(personal_info.get('phone'))
    if personal_info.get('location'): contact_parts.append(personal_info.get('location'))
    if personal_info.get('linkedin'): contact_parts.append(personal_info.get('linkedin'))
    
    if contact_parts:
        doc.add_paragraph(" | ".join(contact_parts))
        
    # Professional Summary
    if optimized_summary:
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(optimized_summary)
        
    # Experience
    if experience:
        doc.add_heading("Professional Experience", level=1)
        for exp in experience:
            job_title = exp.get('jobTitle', '')
            company = exp.get('company', '')
            dates = f"{exp.get('startDate', '')} - {exp.get('endDate', 'Present')}"
            
            p = doc.add_paragraph()
            p.add_run(f"{job_title} at {company}").bold = True
            doc.add_paragraph(dates)
            
            desc = exp.get('description', '')
            if desc:
                doc.add_paragraph(desc)
                
    # Education
    if education:
        doc.add_heading("Education", level=1)
        for ed in education:
            degree = ed.get('degree', '')
            field = ed.get('fieldOfStudy', '')
            institution = ed.get('institution', '')
            year = ed.get('graduationYear', '')
            
            p = doc.add_paragraph()
            p.add_run(f"{degree} in {field}").bold = True
            doc.add_paragraph(f"{institution}, {year}")
            
    # Skills
    if optimized_skills:
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(optimized_skills))
        
    # Save to BytesIO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f"Optimized_Resume_{name.replace(' ', '_')}.docx" if name else "Optimized_Resume.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@builder_bp.route('/review', methods=['POST'])
def review_resume():
    # Force reload of .env to get the latest GEMINI_API_KEY without restarting server
    env_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), '.env')
    load_dotenv(env_path, override=True)

    data = request.get_json()
    if not data:
        return jsonify({"error": "No data provided"}), 400
        
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        return jsonify({"error": "GEMINI_API_KEY is not set in the environment"}), 500
        
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-flash-latest')
        
        prompt = f"""
        You are an expert technical recruiter and resume reviewer.
        Review the following resume data and provide a concise, actionable, and professional critique.
        Highlight what is done well and what needs improvement.
        
        Resume Data:
        {data}
        
        Format the response clearly using Markdown. Keep it under 300 words.
        """
        
        response = model.generate_content(prompt)
        return jsonify({"review": response.text}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500
